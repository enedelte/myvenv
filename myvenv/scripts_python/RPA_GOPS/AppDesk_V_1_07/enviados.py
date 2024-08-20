import win32com.client as win32
import os
import csv
import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def guardar_correos_en_csv(correos, nombre_archivo):
    """
    Guarda los correos en un archivo CSV.
    """
    try:
        with open(nombre_archivo, 'w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            writer.writerow(['Asunto', 'Remitente', 'Fecha de Envío', 'Destinatarios', 'Cuerpo', 'Adjuntos'])
            for correo in correos:
                writer.writerow([
                    correo['Asunto'],
                    correo['Remitente'],
                    correo['Fecha de Envío'],
                    correo['Destinatarios'],
                    correo['Cuerpo'],
                    correo['Adjuntos']
                ])
        logging.info(f"Los correos han sido guardados en el archivo '{nombre_archivo}'.")
    except Exception as e:
        messagebox.showerror("Error", f"Error al guardar los correos en '{nombre_archivo}': {e}")
        logging.error(f"Error al guardar los correos en '{nombre_archivo}': {e}")

def correos_enviados_por_fecha(fecha_deseada):
    """
    Retorna los correos enviados en una fecha específica.
    """
    outlook = win32.Dispatch("Outlook.Application")
    enviados = outlook.GetNamespace("MAPI").GetDefaultFolder(5)
    correos_filtrados = []
    for item in enviados.Items:
        if item.Class == 43 and item.SentOn.date() == fecha_deseada:  # Verificar si es un correo electrónico y si la fecha de envío coincide
            correo = {
                "Asunto": item.Subject,
                "Remitente": item.SenderName,
                "Fecha de Envío": item.SentOn,
                "Destinatarios": item.To,
                "Cuerpo": item.Body,
                "Adjuntos": ';'.join([attachment.FileName for attachment in item.Attachments])
            }
            correos_filtrados.append(correo)
    return correos_filtrados

def leer_correos_desde_csv(nombre_archivo):
    """
    Lee los correos desde un archivo CSV y los devuelve como una lista de diccionarios.
    """
    try:
        correos = []
        with open(nombre_archivo, 'r', newline='', encoding='utf-8') as file:
            reader = csv.reader(file)
            next(reader)  # Saltar la primera fila, que contiene los encabezados
            for row in reader:
                correo = {
                    "Asunto": row[0],
                    "Remitente": row[1],
                    "Fecha de Envío": row[2],
                    "Destinatarios": row[3],
                    "Cuerpo": row[4],
                    "Adjuntos": row[5]
                }
                correos.append(correo)
        return correos
    except Exception as e:
        messagebox.showerror("Error", f"Error al leer los correos desde '{nombre_archivo}': {e}")
        logging.error(f"Error al leer los correos desde '{nombre_archivo}': {e}")
        return []

def escribir_correo_en_docx(correo, nombre_archivo):
    """
    Escribe un correo en un archivo DOCX con un formato específico.
    """
    try:
        document = Document()

        # Establecer estilos de texto
        estilo_titulo = document.styles['Title']
        estilo_titulo.font.size = Pt(12)
        estilo_titulo.font.bold = True

        estilo_normal = document.styles['Normal']
        estilo_normal.font.size = Pt(10)

        # Establecer el espacio entre párrafos
        espacio_entre_parrafos = Pt(8)

        # Añadir título (asunto)
        titulo = document.add_heading(correo["Asunto"], level=1)
        titulo.style = estilo_titulo

        # Añadir párrafos
        remitente = document.add_paragraph("Remitente: " + correo["Remitente"] + " respuestaspqr@nuevaeps.com.co")
        destinatarios = document.add_paragraph("Destinatarios: " + correo["Destinatarios"])
        fecha_envio = document.add_paragraph("Fecha de Envío: " + correo["Fecha de Envío"])
        adjuntos = document.add_paragraph("Adjuntos: " + correo["Adjuntos"])
        cuerpo = document.add_paragraph("Cuerpo: " + correo["Cuerpo"])

        # Aplicar estilos a los párrafos
        for paragraph in [remitente, fecha_envio, destinatarios, cuerpo, adjuntos]:
            paragraph.style = estilo_normal
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alinear a la izquierda
            paragraph.space_after = espacio_entre_parrafos
            paragraph.space_before = espacio_entre_parrafos
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Color de texto negro

        document.save(nombre_archivo)
        logging.info(f"El correo '{correo['Asunto']}' ha sido guardado en el archivo '{nombre_archivo}'.")
    except Exception as e:
        messagebox.showerror("Error", f"Error al escribir el correo en '{nombre_archivo}': {e}")
        logging.error(f"Error al escribir el correo en '{nombre_archivo}': {e}")

def seleccionar_ruta():
    """
    Permite al usuario seleccionar una carpeta de destino para guardar los archivos DOCX.
    """
    # Crear la ventana principal de Tkinter
    root = tk.Tk()
    root.withdraw()  # Ocultar la ventana principal

    # Solicitar la carpeta de destino para guardar los archivos DOCX
    carpeta_destino = filedialog.askdirectory(title="Seleccionar carpeta de destino para archivos DOCX")

    if not carpeta_destino:
        messagebox.showwarning("Advertencia", "No se seleccionó ninguna carpeta.")
        return ""

    return carpeta_destino

def main():
    """
    Función principal del programa.
    """
    try:
        # Obtener la fecha deseada (hoy)
        fecha_deseada = datetime.date.today()

        # Obtener los correos enviados en la fecha deseada
        correos = correos_enviados_por_fecha(fecha_deseada)

        # Guardar los correos en un archivo CSV
        nombre_archivo_csv = 'correos.csv'
        guardar_correos_en_csv(correos, nombre_archivo_csv)

        # Leer los correos desde el archivo CSV
        correos_leidos = leer_correos_desde_csv(nombre_archivo_csv)

        # Solicitar la carpeta de destino para guardar los archivos DOCX
        carpeta_destino = seleccionar_ruta()
        if not carpeta_destino:
            return

        # Guardar cada correo en un archivo DOCX individual
        for correo in correos_leidos:
            nombre_archivo_docx = os.path.join(carpeta_destino, f"{correo['Asunto'][27:34]}.docx")
            escribir_correo_en_docx(correo, nombre_archivo_docx)

        messagebox.showinfo("Completado", "Se han guardado los correos electrónicos como archivos DOCX en la carpeta seleccionada.")
        logging.info("Proceso completado exitosamente.")
    except Exception as e:
        messagebox.showerror("Error", f"Se ha producido un error inesperado: {e}")
        logging.error(f"Error inesperado: {e}")

if __name__ == "__main__":
    main()

